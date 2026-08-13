# -*- coding: utf-8 -*-
"""Build an ATS-friendly .docx CV for Sahatsawat Thongma.
Single column, native text (clean Thai + English), standard headings.
Run: python build_cv_docx.py  ->  cv-sahatsawat-thongma.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1A, 0x1A, 0x1A)
BLUE = RGBColor(0x00, 0x43, 0xCE)
GREY = RGBColor(0x55, 0x55, 0x55)
FONT = "Tahoma"  # Thai-capable, ubiquitous, ATS-safe

doc = Document()

# ---- base style + page margins ----
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
# ensure Thai (complex-script) also uses the same font
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:cs"), FONT)
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)
pf = normal.paragraph_format
pf.space_after = Pt(4)
pf.line_spacing = 1.12

for s in doc.sections:
    s.top_margin = Mm(15)
    s.bottom_margin = Mm(15)
    s.left_margin = Mm(16)
    s.right_margin = Mm(16)


def _set_run_fonts(run):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rf.set(qn("w:cs"), FONT)


def add_run(p, text, bold=False, size=None, color=None):
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    r.font.color.rgb = color if color else INK
    _set_run_fonts(r)
    return r


def para(space_after=4, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    return p


def bottom_border(p):
    """Add a bottom border (rule) under a paragraph -> section underline."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1A1A1A")
    pbdr.append(bottom)
    pPr.append(pbdr)


def section_heading(text):
    p = para(space_before=10, space_after=5)
    add_run(p, text.upper(), bold=True, size=11, color=INK)
    bottom_border(p)


def bullet(runs):
    """runs: list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    for text, bold in runs:
        add_run(p, text, bold=bold)
    return p


# ===================== HEADER =====================
p = para(space_after=1)
add_run(p, "สหัสวรรษ ทองมา (Sahatsawat Thongma)", bold=True, size=20, color=INK)

p = para(space_after=1)
add_run(p, "Data Engineer", bold=True, size=12.5, color=BLUE)

p = para(space_after=4)
add_run(p, "Data platform · Web apps · AI systems — ออกแบบ สร้าง และดูแลเองตั้งแต่ต้นจนจบ", size=10, color=GREY)

p = para(space_after=1)
add_run(p, "sahatsawat.thongma@gmail.com  |  +66 83-691-4296  |  พันท้ายนรสิงห์, สมุทรสาคร (เปิดรับ remote / hybrid)", size=9.5, color=GREY)
p = para(space_after=2)
add_run(p, "github.com/sthongma  |  sthongma.github.io/portfolio/resume.html", size=9.5, color=GREY)

# ===================== SUMMARY =====================
section_heading("Professional Summary")
p = para(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_run(p,
    "Data Engineer ประสบการณ์ 4 ปีในสาย e-commerce — ออกแบบ สร้าง และดูแล data platform "
    "เต็มระบบแบบ end-to-end ให้บริษัทที่ขายผ่าน 4 marketplace ราว 2,000 ออเดอร์/วัน รวมถึง web app และ AI system "
    "ที่ทีมใช้งานจริงทุกวัน งานมือ ~40 งาน/วันเปลี่ยนเป็นระบบอัตโนมัติที่รันเองตั้งแต่ตีห้า "
    "และผู้บริหารดูยอดขาย/สต็อกได้เองโดยไม่ต้องรอรายงาน")
p = para(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_run(p,
    "เทคโนโลยีที่ใช้จริงใน production: PostgreSQL data warehouse 52 GB (70 ล้านแถว, 8 source systems), "
    "dbt (199 models / 494 tests), Apache Airflow ใน Docker, Python, ELT pipelines, "
    "medallion + dimensional (Kimball) modeling, FastAPI/React และ Azure "
    "ปี 2026 ปิด platform migration 2 ตัวจบในไตรมาสเดียว "
    "(Azure SQL Server → PostgreSQL on-prem, Azure Data Factory → Airflow)")
p = para(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_run(p,
    "โฟกัสปัจจุบันคือ data reliability — จับเคสที่ pipeline ขึ้นว่าสำเร็จแต่ข้อมูลหายเงียบ ๆ "
    "ก่อนที่ตัวเลขผิดจะไปถึงการตัดสินใจ")

# ===================== SKILLS =====================
section_heading("Technical Skills")
skills = [
    ("Data Warehouse: ", "PostgreSQL, SQL Server, T-SQL, dbt, Medallion architecture, Kimball / star schema, SCD type 2"),
    ("Data Engineering: ", "Python, Pandas, ETL / ELT, Apache Airflow, Docker, SQLAlchemy"),
    ("Reliability & Observability: ", "dbt tests & data contracts, source freshness SLAs, Elementary Data, pipeline monitoring & alerting, backup & disaster recovery, PostgreSQL administration"),
    ("App Development: ", "FastAPI, React / TypeScript, REST / JSON API"),
    ("AI Systems: ", "LLM tool-calling / function calling, AI agents, natural-language-to-SQL (read-only + PII guard), Azure OpenAI / Azure AI Foundry"),
    ("Ingestion & BI: ", "Playwright (RPA), Power BI, Power Query"),
    ("Cloud — Azure: ", "Blob Storage, App Service, Data Factory (ADF), Azure SQL, Document Intelligence"),
    ("Developer Tooling: ", "Git, GitHub Actions, AI pair-programming (Claude Code, Cursor)"),
]
for label, rest in skills:
    p = para(space_after=3)
    add_run(p, label, bold=True)
    add_run(p, rest)

# ===================== EXPERIENCE =====================
section_heading("Professional Experience")

# --- EP Asia Group ---
p = para(space_after=0)
add_run(p, "EP Asia Group", bold=True, size=11)
add_run(p, "          May 2025 – Present", size=9.5, color=GREY)
p = para(space_after=3)
add_run(p, "Data Engineer — E-Commerce Data Platform", bold=True, size=10, color=BLUE)

bullet([("ทำให้ KPI ที่ผู้บริหารดูทุกวันเชื่อถือได้เป็นครั้งแรก", True),
        (" — อัตราส่งของทันเวลา (fulfillment rate) เคยรายงานว่าปี 2025 หลุดเป้าเกือบทั้งหมด "
         "ทั้งที่ ERP ยืนยันว่าส่งของแล้ว 97% เพราะโมเดลตัดสินช่วงเวลาที่ feed ต้นทางยังไม่เกิด "
         "และเอาออเดอร์ที่สัญญาคนละแบบ (พรีออเดอร์/B2B/ขายส่ง) มารวมในตัวหารเดียวกัน · รื้อวิธีวัดใหม่ทั้งชุด → ปัจจุบัน ", False),
        ("94.6% บน 323,134 ออเดอร์", True),
        (" และเคสที่หลุดซึ่ง ", False), ("88% เคยไม่มีสาเหตุกำกับ", True),
        (" ตอนนี้ระบุสาเหตุได้ครบ ทีมปฏิบัติการไล่แก้เป็นรายสาเหตุได้ครั้งแรก · เทคนิค: data-coverage floor, "
         "scope gate, failure-reason taxonomy 6 ชั้น บน fact table 17 แกนพยาน", False)])

bullet([("จับและแก้ data incident ที่ไม่มีสัญญาณเตือน", True),
        (" — กู้เลขพัสดุคืน ", False), ("23,677 เลข", True),
        (" ที่ dedup rule ลบทิ้งเงียบ ๆ และแก้ที่รากจน ", False),
        ("28.5% ของยอดขายเดือนหนึ่งที่ไม่ติดชื่อร้านเหลือศูนย์", True),
        (" · เทคนิค: freshness canary เทียบ source สด, value-presence gate เป็น hard failure, regression test กันกลับมาโดยไม่มีใครรู้ · "
         "และทำ disaster-recovery backup ตัวแรกของ warehouse (nightly pg_dump ขึ้น Azure Blob พร้อม verification task แยก retention 35 วัน)", False)])

bullet([("เปลี่ยนปัญหาที่เคย “รู้ตอนสายไปแล้ว” ให้ดักได้ล่วงหน้า", True),
        (" — สร้างระบบเทียบราคาขายจริงกับเพดานราคา/ส่วนลดราย SKU × ร้าน (มี.ค. 2026): "
         "ย้อนตรวจทั้งประวัติการขายดักได้ ", False),
        ("150 คู่ SKU×ร้าน / 970 จุด", True),
        (" ที่ให้ส่วนลดเกินเพดานที่ตั้งไว้เอง ", False),
        ("เฉลี่ย 9.9 จุด% (หนักสุดเกิน 180 จุด%) คิดเป็นเงินรั่ว 9.3% ของมูลค่าสินค้าในกลุ่มนั้น", True),
        (" · และสร้าง reorder engine บน dashboard (จุดสั่งซ้ำ, safety stock, days-of-cover, ปริมาณที่ควรสั่ง) "
         "แทนการเปิดหลายตารางมานั่งเทียบเอง — ปัจจุบันครอบ ", False),
        ("2,971 SKU ชี้ 373 ตัวที่ต้องสั่งทันที", True),
        (" และจับ PO ค้างที่ระบบเคยนับผิดว่าเป็น “ของกำลังมา”", False)])

bullet([("แทนที่ vendor WMS ด้วย custom apps ที่เขียนเอง และปิดช่องว่างที่ ERP ทำไม่ได้", True),
        (" — WMS: จากโค้ดบรรทัดแรกถึง production deploy ใน ", False),
        ("15 วัน", True),
        (" พนักงานคลัง (รวมที่ไม่ใช้ภาษาไทย) เรียนรู้ใน 5 นาที ทุกการหยิบ/นับ/ย้ายลง inventory ledger "
         "พร้อม audit trail — ปัจจุบัน ", False),
        ("208,189 รายการ · ผู้ใช้จริง 36 คน", True),
        (" รองรับ ~2,000 ออเดอร์/วัน · ระบบสแกน barcode ปิด 2 ช่องว่างของ ERP: แยก ", False),
        ("“เคลม” ออกจาก “คืนปกติ” ได้ 21,152 พัสดุ", True),
        (" และตอบได้ว่าของกลับถึงคลังหรือยัง (มองเห็นขากลับจาก 26.8% เป็น 98.0% ผ่าน crosswalk "
         "เลขขาไป↔ขากลับ 28,891 คู่) — เห็นของที่กำลังเดินทางกลับอีก ", False),
        ("2,661 พัสดุ", True), (" ที่เดิมมองไม่เห็น", False)])

bullet([("สร้าง AI system ที่ใช้จริงใน production ไม่ใช่แค่ใช้ AI ช่วยเขียนโค้ด", True),
        (" — ผู้ช่วยข้อมูลที่พนักงานถามเป็นภาษาไทยแล้วได้ Excel/dashboard จากคลังข้อมูลจริง "
         "(บังคับ SELECT-only ระดับฐานข้อมูล + guard ข้อมูลส่วนบุคคล), chat widget แบบ function-calling "
         "บน dashboard และบอทสรุปสถานะระบบใน Telegram", False)])

bullet([("ลด cloud cost รายเดือนอย่างวัดได้", True),
        (" — ปิด platform migration 2 ตัวจบในไตรมาสเดียว (Azure SQL Server → self-hosted PostgreSQL "
         "และปลดระวาง Azure Data Factory ทั้ง 46 pipelines → Airflow DAGs) โดย dashboard ไม่ดาวน์เลย · "
         "ตรวจย้อนกับบิล Azure รายเดือนจริง 7 เดือน: ค่าฐานข้อมูลคลังบนคลาวด์ ", False),
        ("โตขึ้น 162% ใน 5 เดือน", True),
        (" (โตเร็วกว่าตัวธุรกิจ) จนกลายเป็นรายจ่ายก้อนใหญ่ที่สุด — หลังย้ายเสร็จ ", False),
        ("ปิดฐานข้อมูลนั้นถาวรเหลือ 0", True),
        (" และ ", False),
        ("ค่า Data Factory เหลือ 0 ตั้งแต่วันถัดจากวันปลดระวาง", True),
        (" · ผลที่ใหญ่กว่าค่าใช้จ่ายคือ ", False),
        ("nightly batch จาก 2 ชม. เหลือ 20 นาที", True),
        (" — เหตุผลที่ย้ายมาจากการวัด ไม่ใช่ความรู้สึก: batch เดิมติด IO 98% บนคลาวด์ "
         "และการเพิ่ม vCore ทดสอบแล้ววัดได้ว่าแพงขึ้นโดยไม่คุ้มจนต้องถอยกลับ · ", False),
        ("ดูแลระบบ production ทั้ง 7 ตัว", True),
        (" ตั้งแต่ architecture, build, deploy ถึง on-call", False)])

# --- Traveler ---
p = para(space_before=8, space_after=0)
add_run(p, "Traveler Co., Ltd.", bold=True, size=11)
add_run(p, "          2022 – 2025", size=9.5, color=GREY)
p = para(space_after=3)
add_run(p, "Accounting · ERP & BI · Operations (ควบหลายหน้าที่)", bold=True, size=10, color=BLUE)

bullet([("จบปัญหาสต็อกไม่ตรงกันข้ามช่องทางขาย", True),
        (" — วาง master data ของ ERP ตั้งแต่ศูนย์: ผังบัญชี, ข้อมูลลูกค้า/ผู้ขาย และมาตรฐาน SKU เดียวกันทุก marketplace ให้ finance กับ sales ใช้ตัวเลขชุดเดียวกัน (Ecount ERP) · เทคนิค: chart of accounts, customer/vendor master, SKU mapping ข้าม marketplaces + ERP", False)])
bullet([("ปิดปัญหาตัวเลข sales กับ finance ไม่ตรงกันที่เกิดซ้ำทุกเดือน", True),
        (" — สร้างรายงานรายเดือนกลางที่สองทีมใช้กระทบยอดร่วมกัน (Power BI, Power Query)", False)])

# ===================== PROJECT =====================
section_heading("Selected Project")
p = para(space_after=2)
add_run(p, "E-Commerce Data Pipeline & Analytics Platform — built & operated end-to-end", bold=True)
p = para(space_after=3)
add_run(p, "52 GB PostgreSQL warehouse · 70M rows · 8 source systems · 199 dbt models · 494 data tests · 11 Airflow DAGs · 32 file types · ~2,000 orders/day", size=9.5, color=GREY)
p = para(space_after=3)
add_run(p, "Data platform เต็มระบบ ออกแบบ สร้าง และดูแลเอง — เปลี่ยนงานที่เคยต้องมีคนนั่งโหลดและ refresh ไฟล์ทุกเช้า "
           "ให้ระบบดึง raw data จาก 8 source systems (marketplaces, OMS, internal apps, logistics) เองอัตโนมัติ "
           "แล้วเสิร์ฟเป็น dashboard แบบ real-time ให้ทั้งบริษัท พร้อมผู้ช่วย AI ให้ถามข้อมูลเป็นภาษาคนได้ "
           "(Python + Playwright → Azure Blob data lake → dbt medallion 5 ชั้นบน PostgreSQL → Apache Airflow ใน Docker → FastAPI + React)")
p = para()
add_run(p, "แดชบอร์ด 5 ตัวที่ใช้งานจริง: KPI การส่งของ (fulfillment) · ดักราคา/ส่วนลดตั้งผิด · "
           "วางแผนสั่งซื้อ (reorder engine) · วิเคราะห์การขาย (ออนไลน์ + ขายส่ง) · ผู้ดูแลระบบและสิทธิ์ผู้ใช้",
        size=9.5, color=GREY)

# ===================== EDUCATION =====================
section_heading("Education")
p = para(space_after=0, space_before=2)
add_run(p, "วิศวกรรมศาสตรบัณฑิต (วศ.บ.) สาขาวิศวกรรมคอมพิวเตอร์ — กำลังศึกษา", bold=True)
p = para(space_after=4)
add_run(p, "มหาวิทยาลัยเอเชียอาคเนย์ (Southeast Asia University) · ระบบการศึกษาทางไกลทางอินเทอร์เน็ต · ส.ค. 2026 – คาดว่าจบปี 2029 · เรียนคู่กับงานประจำ (เข้าเรียนสดช่วงเย็นวันธรรมดาและวันหยุด)", size=9.5, color=GREY)
p = para(space_after=0, space_before=2)
add_run(p, "มัธยมศึกษาตอนปลาย (ม.6) — สายวิทย์-คณิต (Science–Math)", bold=True)
p = para(space_after=2)
add_run(p, "โรงเรียนเตรียมอุดมศึกษาพัฒนาการ อุดรธานี (Triam Udom Suksa Pattanakan Udon Thani) · จบปี 2562 / 2019", size=9.5, color=GREY)

# ===================== CERTIFICATIONS =====================
section_heading("Certifications")
edu = [
    ("Power BI & Excel for Data Analysis",
     "9Expert Online Training — Power BI modeling, DAX และ Power Query (ใช้จริงกับ finance & sales reconciliation ที่ Traveler Co.)"),
    ("Python Programming Foundations",
     "Skooldio · Jun 2025 — Core Python สำหรับงานข้อมูล ใช้เป็นฐานของ ETL scripts, FastAPI services และ Airflow operators ใน production"),
]
for title, meta in edu:
    p = para(space_after=0, space_before=2)
    add_run(p, title, bold=True)
    p = para(space_after=2)
    add_run(p, meta, size=9.5, color=GREY)

doc.save("cv-sahatsawat-thongma.docx")
print("saved cv-sahatsawat-thongma.docx")
